import os
import csv
import time
from docx import Document
from googletrans import Translator
import json

# Initialize the translator
translator = Translator()

# Load the CSV file
csv_file = 'The-Office-Lines-V4.csv'

# Create a directory to store all seasons
base_dir = 'The_Office_Scripts'
if not os.path.exists(base_dir):
    os.mkdir(base_dir)

# Translate text to Turkish with retry mechanism
def translate_to_turkish_batch(texts, max_retries=3):
    for attempt in range(max_retries):
        try:
            translations = translator.translate(texts, dest='tr')
            return [translation.text for translation in translations]
        except Exception as e:
            print(f"Translation error: {e}. Retrying in 5 seconds...")
            time.sleep(5)
    return [f"[Translation failed after {max_retries} attempts: {text}]" for text in texts]

# Load progress
progress_file = 'translation_progress.json'
if os.path.exists(progress_file):
    with open(progress_file, 'r') as f:
        progress = json.load(f)
else:
    progress = {'last_season': 0, 'last_episode': 0}

# Open and read the CSV file
with open(csv_file, 'r', encoding='utf-8') as file:
    csv_reader = csv.reader(file)
    headers = next(csv_reader)  # Read the header row
    print(f"Number of columns in CSV: {len(headers)}")
    print("Column headers:", headers)

    # Determine indices for required columns
    required_columns = ['season', 'episode', 'title', 'speaker', 'line']
    column_indices = {col: headers.index(col) if col in headers else -1 for col in required_columns}

    current_season = None
    current_episode = None
    current_document = None
    lines_to_translate = []
    paragraphs = []

    # Iterate through rows in the CSV file
    for row in csv_reader:
        season = int(row[column_indices['season']]) if column_indices['season'] != -1 else 0
        episode = int(row[column_indices['episode']]) if column_indices['episode'] != -1 else 0

        # Skip already processed episodes
        if season < progress['last_season'] or (season == progress['last_season'] and episode < progress['last_episode']):
            continue

        # Check if we're starting a new episode
        if season != current_season or episode != current_episode:
            # Translate and save the previous episode if it exists
            if current_document:
                translated_lines = translate_to_turkish_batch(lines_to_translate)
                for i, (speaker, line) in enumerate(paragraphs):
                    # Add original line
                    p = current_document.add_paragraph()
                    p.add_run(f'{speaker}: ').bold = True
                    p.add_run(line)

                    # Add translated line
                    p = current_document.add_paragraph()
                    p.add_run(translated_lines[i])
                    p.style = 'Quote'

                    # Add a blank line for readability
                    current_document.add_paragraph()

                episode_dir = os.path.join(base_dir, f'Season_{current_season}')
                if not os.path.exists(episode_dir):
                    os.makedirs(episode_dir)

                current_document.save(os.path.join(episode_dir, f'Episode_{current_episode}.docx'))
                print(f"Saved: Season {current_season}, Episode {current_episode}, Lines {len(lines_to_translate)}")

            # Start a new document for the new episode
            current_season = season
            current_episode = episode
            current_document = Document()
            title = row[column_indices['title']] if column_indices['title'] != -1 else 'Unknown'
            current_document.add_heading(f'Season {season}, Episode {episode}: {title}', level=1)
            lines_to_translate = []
            paragraphs = []

        speaker = row[column_indices['speaker']] if column_indices['speaker'] != -1 else 'Unknown'
        line = row[column_indices['line']] if column_indices['line'] != -1 else 'No line'

        # Collect lines for batch translation
        lines_to_translate.append(line)
        paragraphs.append((speaker, line))  # Store speaker and line for later use

    # Translate and save the last episode
    if current_document:
        translated_lines = translate_to_turkish_batch(lines_to_translate)
        for i, (speaker, line) in enumerate(paragraphs):
            # Add original line
            p = current_document.add_paragraph()
            p.add_run(f'{speaker}: ').bold = True
            p.add_run(line)

            # Add translated line
            p = current_document.add_paragraph()
            p.add_run(translated_lines[i])
            p.style = 'Quote'

            # Add a blank line for readability
            current_document.add_paragraph()

        episode_dir = os.path.join(base_dir, f'Season_{current_season}')
        if not os.path.exists(episode_dir):
            os.makedirs(episode_dir)

        current_document.save(os.path.join(episode_dir, f'Episode_{current_episode}.docx'))
        print(f"Saved: Season {current_season}, Episode {current_episode}, Lines {len(lines_to_translate)}")

    # Update progress after processing all episodes
    progress['last_season'] = current_season
    progress['last_episode'] = current_episode
    with open(progress_file, 'w') as f:
        json.dump(progress, f)

print("All episodes have been processed and saved.")
